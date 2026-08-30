"""Load real user interviews from a Word document into research records."""

from __future__ import annotations

import hashlib
from pathlib import Path

from common.models import SourceType
from ingestion.connectors.research_excel import AGE_25_35
from ingestion.filters.relevance import is_relevant
from ingestion.schemas import RawRecord

DEFAULT_INTERVIEW_DOCX = "All file.docx"


def _author_hash(key: str) -> str:
    return hashlib.sha256(key.encode()).hexdigest()[:32]


def _load_paragraphs(path: Path) -> list[str]:
    from docx import Document

    document = Document(str(path))
    return [para.text.strip() for para in document.paragraphs if para.text.strip()]


def _load_tables_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" — ".join(cells))
    return "\n".join(lines)


def _slice(paragraphs: list[str], start: str, end: str | None) -> str:
    start_i = next((i for i, text in enumerate(paragraphs) if start in text), None)
    if start_i is None:
        return ""
    if end is None:
        return "\n".join(paragraphs[start_i:])
    end_i = next((i for i, text in enumerate(paragraphs[start_i + 1 :], start_i + 1) if end in text), len(paragraphs))
    return "\n".join(paragraphs[start_i:end_i])


def fetch_interview_docx_records(docx_path: str | Path | None = None) -> list[RawRecord]:
    """Parse All file.docx into one research document per interview (age 25–35)."""
    path = Path(docx_path or DEFAULT_INTERVIEW_DOCX)
    if not path.exists():
        return []

    paragraphs = _load_paragraphs(path)
    interviews = [
        (
            "snehi",
            "Snehi",
            _slice(paragraphs, "How long have you been using Myntra?", "Candidate [Snehi]"),
        ),
        (
            "dresses-weekly",
            "Interview 2",
            _slice(
                paragraphs,
                "How long have you used Myntra, and what do you usually shop for? 2015",
                "How long have you used Myntra, and what do you usually shop for? 5 years",
            ),
        ),
        (
            "laptop-bag",
            "Interview 3",
            _slice(paragraphs, "How long have you used Myntra, and what do you usually shop for? 5 years", "A month back"),
        ),
        (
            "workout",
            "Interview 4",
            _slice(paragraphs, "A month back, workut clothes", "Interview Guide — Candidate 1"),
        ),
        (
            "candidate-1-final",
            "Interview 5",
            _slice(paragraphs, "Interview Guide — Candidate 1", "Backup probe bank"),
        ),
        (
            "gurgaon-notes",
            "Interview 6",
            _load_tables_text(path),
        ),
    ]

    records: list[RawRecord] = []
    for index, (slug, label, body) in enumerate(interviews):
        text = " ".join(body.split())
        if len(text) < 80:
            continue
        prefixed = f"Age band: 25-35. Real interview ({label}). {text}"
        relevance = is_relevant(prefixed, always_include=True)
        records.append(
            RawRecord(
                source=SourceType.research,
                source_ref=f"research:interview-docx:row:{index}",
                text=prefixed,
                author_hash=_author_hash(f"interview-docx:{slug}"),
                language="en",
                metadata={
                    "sheet": path.name,
                    "row_index": index,
                    "age_band": AGE_25_35,
                    "workbook": "interview-docx",
                    "interview_id": slug,
                    "kind": "interview",
                },
                matched_signals=list(relevance.matched_signals),
            )
        )
    return records
